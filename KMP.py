import docx
import time
import random


def KMPSearch(pat, txt):
    M = len(pat)
    N = len(txt)

    # create lps[] that will hold the longest prefix suffix
    # values for pattern
    lps = [0]*M
    j = 0 # index for pat[]

    # Preprocess the pattern (calculate lps[] array)
    computeLPSArray(pat, M, lps) # calculate longest prefix suffix values are stored in lps

    i = 0 # index for txt[]
    indices = [] # initialize list of indices
    while i < N:
        if pat[j] == txt[i]:
            i += 1 # proceed to the next characters 
            j += 1

        if j == M: #complete match is found 
            indices.append(i-j) # pattern found, add index to list
            j = lps[j-1]

        # mismatch after j matches
        elif i < N and pat[j] != txt[i]: #avoid unnecesary comparisons by skipping ahead
            # Do not match lps[0..lps[j-1]] characters,
            # they will match anyway
            if j != 0:
                j = lps[j-1]
            else:
                i += 1

    return indices # return list of indices where pattern is found




def computeLPSArray(pat, M, lps):
    len = 0 # length of the previous longest prefix suffix

    lps[0] = 0
    i = 1 #second index of the pattern, first index doesnT have any proper suffixes

    # the loop calculates lps[i] for i = 1 to M-1
    while i < M:
        if pat[i] == pat[len]: #means we found a longer proper suffix, meaning the suffix that is not equal to the entire string itself
            len += 1 # upgrade the length of the suffix 
            lps[i] = len
            i += 1 # next position
        else: #current suffix is not a proper suffix anymore
            # This is tricky. Consider the example.
            # AAACAAAA and i = 7. The idea is similar
            # to search step.
            if len != 0:
                len = lps[len-1]

                # Also, note that we do not increment i here
            else:
                lps[i] = 0
                i += 1





# Open the Word file
#doc = docx.Document(r"C:\Users\lenovo\Downloads\Raven.docx")

#text = " "

# Iterate over paragraphs in the document
#for para in doc.paragraphs:
 #   # Print the text of each paragraph
  #  text += para.text

doc = docx.Document(r"C:\Users\lenovo\Downloads\Raven.docx")
text = '\n'.join([para.text for para in doc.paragraphs])


#small pattern
pat_small = "silence"
start_time = time.time()
indices_small = KMPSearch(pat_small,text)
elapsed_time_small = time.time() - start_time
print(f"Indices of '{pat_small}' found using KMP algorithm: {indices_small}") #it says none
print(f"Elapsed time for small pattern: {elapsed_time_small:.6f} seconds")

#large pattern
pat_large = "Ah, distinctly I remember it was in the bleak December"
start_time = time.time()
indices_large = KMPSearch(pat_large,text)
elapsed_time_large = time.time() - start_time
print(f"Indices of large pattern found using KMP algorithm: {indices_large}")
print(f"Elapsed time for large pattern: {elapsed_time_large:.6f} seconds")








import docx
import time
import random


def KMPSearch(pat, txt):
    M = len(pat)
    N = len(txt)

    # create lps[] that will hold the longest prefix suffix
    # values for pattern
    lps = [0]*M
    j = 0 # index for pat[]

    # Preprocess the pattern (calculate lps[] array)
    computeLPSArray(pat, M, lps)

    i = 0 # index for txt[]
    indices = [] # initialize list of indices
    while i < N:
        if pat[j] == txt[i]:
            i += 1
            j += 1

        if j == M:
            indices.append(i-j) # pattern found, add index to list
            j = lps[j-1]

        # mismatch after j matches
        elif i < N and pat[j] != txt[i]:
            # Do not match lps[0..lps[j-1]] characters,
            # they will match anyway
            if j != 0:
                j = lps[j-1]
            else:
                i += 1

    return indices # return list of indices where pattern is found




def computeLPSArray(pat, M, lps):
    len = 0 # length of the previous longest prefix suffix

    lps[0] = 0
    i = 1

    # the loop calculates lps[i] for i = 1 to M-1
    while i < M:
        if pat[i] == pat[len]:
            len += 1
            lps[i] = len
            i += 1
        else:
            # This is tricky. Consider the example.
            # AAACAAAA and i = 7. The idea is similar
            # to search step.
            if len != 0:
                len = lps[len-1]

                # Also, note that we do not increment i here
            else:
                lps[i] = 0
                i += 1





# Open the Word file
#doc = docx.Document(r"C:\Users\lenovo\Downloads\Raven.docx")

#text = " "

# Iterate over paragraphs in the document
#for para in doc.paragraphs:
 #   # Print the text of each paragraph
  #  text += para.text

doc = docx.Document(r"C:\Users\lenovo\Downloads\Raven.docx")
text = '\n'.join([para.text for para in doc.paragraphs])


#small pattern
pat_small = "silence"
start_time = time.time()
indices_small = KMPSearch(pat_small,text)
elapsed_time_small = time.time() - start_time
print(f"Indices of '{pat_small}' found using KMP algorithm: {indices_small}") #it says none
print(f"Elapsed time for small pattern: {elapsed_time_small:.6f} seconds")

#large pattern
pat_large = "Ah, distinctly I remember it was in the bleak December"
start_time = time.time()
indices_large = KMPSearch(pat_large,text)
elapsed_time_large = time.time() - start_time
print(f"Indices of large pattern found using KMP algorithm: {indices_large}")
print(f"Elapsed time for large pattern: {elapsed_time_large:.6f} seconds")








